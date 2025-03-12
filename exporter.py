import csv
import fitz
import openpyxl
from docx import Document
from docx.shared import Pt

class Exporter:
    def __init__(self, settings):
        self.settings = settings

    def export(self, text, file_path):
        """Export text to the specified file format based on the file extension."""
        if file_path.endswith('.txt'):
            self.export_to_txt(text, file_path)
        elif file_path.endswith('.docx'):
            self.export_to_docx(text, file_path)
        elif file_path.endswith('.html'):
            self.export_to_html(text, file_path)
        elif file_path.endswith('.pdf'):
            self.export_to_pdf(text, file_path)
        elif file_path.endswith('.md'):
            self.export_to_markdown(text, file_path)
        elif file_path.endswith('.rtf'):
            self.export_to_rtf(text, file_path)
        elif file_path.endswith('.csv'):
            self.export_to_csv(text, file_path)
        elif file_path.endswith('.xlsx'):
            self.export_to_excel(text, file_path)
        else:
            raise ValueError("Неподдерживаемый формат файла.")

    @staticmethod
    def get_supported_filetypes():
        """Return a list of supported file types for file dialogs."""
        return [
            ("Text files", "*.txt"),
            ("Word Document", "*.docx"),
            ("HTML files", "*.html"),
            ("PDF files", "*.pdf"),
            ("Markdown files", "*.md"),
            ("RTF files", "*.rtf"),
            ("CSV files", "*.csv"),
            ("Excel files", "*.xlsx"),
        ]

    @staticmethod
    def export_to_txt(text, file_path):
        """Export text to a plain text file."""
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(text)

    def export_to_docx(self, text, file_path):
        """Export text to a DOCX file with formatting from settings."""
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = self.settings.font_family
        run.font.size = Pt(self.settings.font_size)
        doc.save(file_path)

    def export_to_html(self, text, file_path):
        """Export text to an HTML file with CSS styling."""
        html_content = f"""
        <html>
        <head>
        <style>
        body {{ font-family: '{self.settings.font_family}', sans-serif; font-size: {self.settings.font_size}pt; }}
        </style>
        </head>
        <body>
        <p>{text.replace('\n', '<br>')}</p>
        </body>
        </html>
        """
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

    def export_to_pdf(self, text, file_path):
        """Export text to a PDF file with formatting from settings."""
        pdf = fitz.open()
        page = pdf.new_page(width=595, height=842)  # A4 size in points
        rect = fitz.Rect(72, 72, 595 - 72, 842 - 72)  # Margins
        text_settings = {
            'fontsize': self.settings.font_size,
            'fontname': 'helv',  # Helvetica as default
        }
        page.insert_textbox(rect, text, **text_settings)
        pdf.save(file_path)
        pdf.close()

    @staticmethod
    def export_to_markdown(text, file_path):
        """Export text to a Markdown file."""
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(text)

    def export_to_rtf(self, text, file_path):
        """Export text to an RTF file with font formatting."""
        font_family = self.settings.font_family
        font_size = self.settings.font_size * 2  # RTF uses half-points
        rtf_content = r"{\rtf1\ansi\deff0{\fonttbl{\f0 " + font_family + r";}}\f0\fs" + str(font_size) + r" " + text + r"}"
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(rtf_content)

    @staticmethod
    def export_to_csv(text, file_path):
        """Export text to a CSV file, treating each line as a row."""
        lines = text.strip().split('\n')
        with open(file_path, 'w', encoding='utf-8', newline='') as f:
            writer = csv.writer(f)
            for line in lines:
                writer.writerow([line])

    @staticmethod
    def export_to_excel(text, file_path):
        """Export text to an Excel file, treating each line as a row."""
        wb = openpyxl.Workbook()
        ws = wb.active
        lines = text.strip().split('\n')
        for idx, line in enumerate(lines, 1):
            ws.cell(row=idx, column=1, value=line)
        wb.save(file_path)